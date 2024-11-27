import pandas as pd
import time
from docx import Document

def gerar_documentos(modelo_docx, excel_path, output_dir):
    # Carregar o Excel
    df = pd.read_excel(excel_path)
    
    # Carregar o modelo do documento
    modelo = Document(modelo_docx)
    
    for index, row in df.iterrows():
        # Criar uma cópia do modelo para edição
        novo_doc = Document(modelo_docx)
        
        # Iterar sobre os parágrafos e substituir as tags dinamicamente
        for paragrafo in novo_doc.paragraphs:
            for tag, valor in row.items():  # Itera sobre as colunas do Excel
                if f"[{tag}]" in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(f"[{tag}]", str(valor))
        time.sleep(4)
        # Salvar o documento personalizado
        output_path = f"{output_dir}/documento_{row['Nome']}.docx"
        novo_doc.save(output_path)
        print(f"Documento salvo: {output_path}")
        time.sleep(4)
# Caminhos dos arquivos
modelo_docx = "modelo.docx"  # Caminho do modelo
excel_path = "dados.xlsx"    # Caminho do Excel
output_dir = "docs_gerados"  # Diretório de saída

# Chamar a função
gerar_documentos(modelo_docx, excel_path, output_dir)
